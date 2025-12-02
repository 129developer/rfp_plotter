"""
Document parsing utilities for extracting text from various file formats.
Supports PDF, DOCX, and plain text files.
"""

import os
from typing import Optional, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Utility class for parsing different document formats"""
    
    @staticmethod
    def extract_text_from_file(file_path: str) -> str:
        """
        Extract text content from a file based on its extension.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return DocumentParser._extract_from_pdf(file_path)
        elif extension == '.docx':
            return DocumentParser._extract_from_docx(file_path)
        elif extension in ['.txt', '.md']:
            return DocumentParser._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    @staticmethod
    def _extract_from_pdf(file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            
            return "\n\n".join(text_content)
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            raise
    
    @staticmethod
    def _extract_from_docx(file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("--- Table ---\n" + "\n".join(table_text))
            
            return "\n\n".join(text_content)
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise
    
    @staticmethod
    def _extract_from_text(file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            raise
    
    @staticmethod
    def get_document_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        stat = file_path.stat()
        
        metadata = {
            'filename': file_path.name,
            'extension': file_path.suffix.lower(),
            'size_bytes': stat.st_size,
            'created_time': stat.st_ctime,
            'modified_time': stat.st_mtime,
        }
        
        # Add format-specific metadata
        if file_path.suffix.lower() == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        'page_count': len(pdf_reader.pages),
                        'pdf_metadata': pdf_reader.metadata if pdf_reader.metadata else {}
                    })
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata: {e}")
        
        elif file_path.suffix.lower() == '.docx':
            try:
                doc = Document(file_path)
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                })
                
                # Try to get document properties
                if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                    metadata['title'] = doc.core_properties.title
                if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                    metadata['author'] = doc.core_properties.author
                if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                    metadata['document_created'] = doc.core_properties.created
                    
            except Exception as e:
                logger.warning(f"Could not extract DOCX metadata: {e}")
        
        return metadata


def validate_document_file(file_path: str) -> bool:
    """
    Validate if a file exists and is a supported document format.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        True if file is valid and supported, False otherwise
    """
    if not os.path.exists(file_path):
        return False
    
    supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
    file_extension = Path(file_path).suffix.lower()
    
    return file_extension in supported_extensions